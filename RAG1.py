"""
RAG 文档处理 - 文档读取与数据清洗模块
负责多种格式文档的读取、解析和数据清洗
"""

import os
import re
import unicodedata
import html
from io import BytesIO
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass, field
import pdfplumber
import pandas as pd
from docx import Document as DocxDocument
from openpyxl import load_workbook


# ==================== 数据结构 ====================

@dataclass
class DocumentPage:
    """文档页面数据结构"""
    page_num: int
    text: str
    tables: List[List[List]] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass
class ParsedDocument:
    """解析后的文档"""
    file_path: str
    file_name: str
    file_type: str
    pages: List[DocumentPage]
    total_pages: int
    metadata: Dict[str, Any] = field(default_factory=dict)


# ==================== 数据清洗器 ====================

class DataCleaner:
    """数据清洗器 - 处理各种文档类型的文本清洗"""

    def __init__(self, config: Optional[Dict] = None):
        self.config = {
            'remove_empty_lines': True,
            'remove_special_chars': True,
            'normalize_whitespace': True,
            'min_line_length': 2,
            'max_line_length': None,
            'remove_urls': False,
            'remove_emails': False,
            'remove_numbers': False,
            'language': 'auto',
            **(config or {})
        }

        # 特殊字符模式
        self.special_chars_pattern = re.compile(
            r'[^\u4e00-\u9fff\u3400-\u4dbf\w\s\.\,\!\?\;\:\'\"\(\)\[\]\{\}\<\>\/\\\|\-\=\+\*\&\^\$\#\@\~`]'
        )

        # URL 模式
        self.url_pattern = re.compile(
            r'https?://(?:[-\w.]|(?:%[\da-fA-F]{2}))+[^\s]*|'
            r'www\.[-\w.]+[^\s]*'
        )

        # 邮箱模式
        self.email_pattern = re.compile(
            r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        )

        # 噪声模式
        self.noise_patterns = [
            re.compile(r'^\s*\d+\s*$'),
            re.compile(r'^\s*第\s*\d+\s*页\s*$'),
            re.compile(r'^\s*Page\s+\d+\s*$', re.IGNORECASE),
            re.compile(r'^\s*\d+\s*/\s*\d+\s*$'),
            re.compile(r'^\s*Copyright\s+©?\s*\d{4}\s*$', re.IGNORECASE),
            re.compile(r'^\s*All\s+Rights\s+Reserved\s*$', re.IGNORECASE),
            re.compile(r'^\s*Confidential\s*$', re.IGNORECASE),
            re.compile(r'^[-=_*]{10,}$'),
            re.compile(r'^[─━]{10,}$'),
        ]

    def clean_text(self, text: str, verbose: bool = False) -> str:
        """主清洗函数"""
        if not text:
            return ""

        original_length = len(text)

        # 1. 标准化 Unicode
        text = self._normalize_unicode(text)

        # 2. 移除控制字符
        text = self._remove_control_chars(text)

        # 3. 标准化空白字符
        if self.config['normalize_whitespace']:
            text = self._normalize_whitespace(text)

        # 4. 移除URL
        if self.config['remove_urls']:
            text = self._remove_urls(text)

        # 5. 移除邮箱
        if self.config['remove_emails']:
            text = self._remove_emails(text)

        # 6. 移除特殊字符
        if self.config['remove_special_chars']:
            text = self._remove_special_chars(text)

        # 7. 移除数字
        if self.config['remove_numbers']:
            text = self._remove_numbers(text)

        # 8. 按行清理
        lines = text.split('\n')
        cleaned_lines = []

        for line in lines:
            cleaned_line = self._clean_line(line)
            if cleaned_line:
                cleaned_lines.append(cleaned_line)

        cleaned_text = '\n'.join(cleaned_lines)

        if verbose:
            reduction = original_length - len(cleaned_text)
            print(f"  清洗统计: {original_length} -> {len(cleaned_text)} 字符 (减少 {reduction})")

        return cleaned_text

    def _normalize_unicode(self, text: str) -> str:
        text = unicodedata.normalize('NFKC', text)
        text = text.replace('\u3000', ' ')
        return text

    def _remove_control_chars(self, text: str) -> str:
        return re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)

    def _normalize_whitespace(self, text: str) -> str:
        text = re.sub(r'[ \t]+', ' ', text)
        lines = [line.strip() for line in text.split('\n')]
        return '\n'.join(lines)

    def _remove_urls(self, text: str) -> str:
        return self.url_pattern.sub('', text)

    def _remove_emails(self, text: str) -> str:
        return self.email_pattern.sub('', text)

    def _remove_special_chars(self, text: str) -> str:
        return self.special_chars_pattern.sub('', text)

    def _remove_numbers(self, text: str) -> str:
        return re.sub(r'\b\d+(?:\.\d+)?\b', '', text)

    def _clean_line(self, line: str) -> Optional[str]:
        line = line.strip()

        if not line:
            return None if self.config['remove_empty_lines'] else ''

        if len(line) < self.config['min_line_length']:
            return None

        if self.config['max_line_length'] and len(line) > self.config['max_line_length']:
            line = line[:self.config['max_line_length']]

        for pattern in self.noise_patterns:
            if pattern.match(line):
                return None

        return line


class TableCleaner:
    """表格数据清洗器"""

    @staticmethod
    def clean_table_data(table_data: List[List]) -> List[List]:
        """清洗表格数据"""
        if not table_data:
            return []

        cleaned_table = []
        for row in table_data:
            cleaned_row = []
            for cell in row:
                if cell is None:
                    cleaned_row.append('')
                else:
                    cleaned = str(cell).strip()
                    cleaned = re.sub(r'\s+', ' ', cleaned)
                    cleaned_row.append(cleaned)
            cleaned_table.append(cleaned_row)

        return cleaned_table

    @staticmethod
    def table_to_markdown(table_data: List[List]) -> str:
        """将表格转换为 Markdown 格式"""
        if not table_data or len(table_data) < 2:
            return ""

        lines = []
        header = "| " + " | ".join(str(cell) for cell in table_data[0]) + " |"
        lines.append(header)
        separator = "| " + " | ".join(["---"] * len(table_data[0])) + " |"
        lines.append(separator)

        for row in table_data[1:]:
            line = "| " + " | ".join(str(cell) for cell in row) + " |"
            lines.append(line)

        return "\n".join(lines)

    @staticmethod
    def table_to_text(table_data: List[List]) -> str:
        """将表格转换为文本格式"""
        if not table_data:
            return ""

        result = []
        for i, row in enumerate(table_data):
            if i == 0:
                # 表头
                result.append(" | ".join(str(cell) for cell in row))
                result.append(" | ".join(["---"] * len(row)))
            else:
                result.append(" | ".join(str(cell) for cell in row))

        return "\n".join(result)


class NoiseFilter:
    """噪声内容过滤器"""

    def __init__(self):
        self.noise_keywords = [
            '请勿转载', '版权所有', '翻印必究',
            'confidential', 'proprietary', 'all rights reserved',
            '仅供内部使用', '内部资料', '注意保密',
        ]

    def is_noise_line(self, line: str, threshold: float = 0.3) -> bool:
        """判断是否为噪声行"""
        if not line:
            return True

        line_lower = line.lower()
        noise_count = sum(1 for kw in self.noise_keywords if kw.lower() in line_lower)

        if noise_count / max(len(line) / 20, 1) > threshold:
            return True

        if re.match(r'^[\W_]+$', line):
            return True

        return False

    def filter_text(self, text: str) -> str:
        """过滤文本中的噪声行"""
        lines = text.split('\n')
        filtered_lines = [line for line in lines if not self.is_noise_line(line)]
        return '\n'.join(filtered_lines)


# ==================== 文档解析器 ====================

class PDFParser:
    """PDF 文档解析器"""

    def __init__(self, cleaner: DataCleaner = None, table_cleaner: TableCleaner = None):
        self.cleaner = cleaner or DataCleaner()
        self.table_cleaner = table_cleaner or TableCleaner()

    def parse(
            self,
            file_path: str,
            from_page: int = 0,
            to_page: int = 100000,
            enable_cleaning: bool = True,
            extract_tables: bool = True,
            verbose: bool = False
    ) -> ParsedDocument:
        """
        解析 PDF 文件

        Args:
            file_path: PDF 文件路径
            from_page: 起始页
            to_page: 结束页
            enable_cleaning: 是否启用清洗
            extract_tables: 是否提取表格
            verbose: 是否打印详细信息

        Returns:
            ParsedDocument 对象
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        pages = []

        with pdfplumber.open(file_path) as pdf:
            total_pages = len(pdf.pages)
            start_page = max(0, from_page)
            end_page = min(total_pages, to_page)

            if verbose:
                print(f"解析 PDF: {file_path}")
                print(f"总页数: {total_pages}, 解析范围: {start_page + 1}-{end_page}")

            for page_num in range(start_page, end_page):
                page = pdf.pages[page_num]

                # 提取文本
                text = page.extract_text() or ""
                if enable_cleaning and text:
                    text = self.cleaner.clean_text(text)

                # 提取表格
                tables = []
                if extract_tables:
                    raw_tables = page.extract_tables()
                    for table in raw_tables:
                        if table and len(table) > 0:
                            if enable_cleaning:
                                cleaned = self.table_cleaner.clean_table_data(table)
                                tables.append(cleaned)
                            else:
                                tables.append(table)

                pages.append(DocumentPage(
                    page_num=page_num + 1,
                    text=text,
                    tables=tables
                ))

                if verbose and (page_num + 1) % 10 == 0:
                    print(f"  已解析 {page_num + 1}/{end_page} 页")

        return ParsedDocument(
            file_path=file_path,
            file_name=os.path.basename(file_path),
            file_type='pdf',
            pages=pages,
            total_pages=len(pages)
        )

    def parse_simple(self, file_path: str, from_page: int = 0, to_page: int = 100000) -> str:
        """简单解析，只返回纯文本"""
        doc = self.parse(file_path, from_page, to_page, enable_cleaning=False, extract_tables=False)
        return "\n\n".join([page.text for page in doc.pages if page.text])


class DOCXParser:
    """DOCX 文档解析器"""

    def __init__(self, cleaner: DataCleaner = None):
        self.cleaner = cleaner or DataCleaner()

    def parse(self, file_path: str, enable_cleaning: bool = True, verbose: bool = False) -> ParsedDocument:
        """解析 DOCX 文件"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        doc = DocxDocument(file_path)
        pages = []
        current_text = []
        page_num = 1

        if verbose:
            print(f"解析 DOCX: {file_path}")

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                if enable_cleaning:
                    text = self.cleaner.clean_text(text)
                if text:
                    current_text.append(text)

            # 检查分页（DOCX 中通过分页符判断）
            if paragraph.runs:
                for run in paragraph.runs:
                    if 'lastRenderedPageBreak' in run._element.xml:
                        if current_text:
                            pages.append(DocumentPage(
                                page_num=page_num,
                                text='\n'.join(current_text),
                                tables=[]
                            ))
                            current_text = []
                            page_num += 1

        # 添加最后一页
        if current_text:
            pages.append(DocumentPage(
                page_num=page_num,
                text='\n'.join(current_text),
                tables=[]
            ))

        return ParsedDocument(
            file_path=file_path,
            file_name=os.path.basename(file_path),
            file_type='docx',
            pages=pages,
            total_pages=len(pages)
        )

    def parse_simple(self, file_path: str) -> str:
        """简单解析，只返回纯文本"""
        doc = self.parse(file_path, enable_cleaning=False)
        return "\n\n".join([page.text for page in doc.pages if page.text])


class TXTParser:
    """TXT 文档解析器"""

    def __init__(self, cleaner: DataCleaner = None):
        self.cleaner = cleaner or DataCleaner()

    def parse(self, file_path: str, encoding: str = 'utf-8', enable_cleaning: bool = True) -> ParsedDocument:
        """解析 TXT 文件"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        # 尝试多种编码
        encodings = [encoding, 'utf-8', 'gbk', 'gb2312', 'ascii']
        content = None

        for enc in encodings:
            try:
                with open(file_path, 'r', encoding=enc) as f:
                    content = f.read()
                break
            except UnicodeDecodeError:
                continue

        if content is None:
            raise ValueError(f"无法解码文件: {file_path}")

        if enable_cleaning:
            content = self.cleaner.clean_text(content)

        pages = [DocumentPage(
            page_num=1,
            text=content,
            tables=[]
        )]

        return ParsedDocument(
            file_path=file_path,
            file_name=os.path.basename(file_path),
            file_type='txt',
            pages=pages,
            total_pages=1
        )

    def parse_simple(self, file_path: str, encoding: str = 'utf-8') -> str:
        """简单解析，只返回纯文本"""
        doc = self.parse(file_path, encoding, enable_cleaning=False)
        return doc.pages[0].text if doc.pages else ""


class ExcelParser:
    """Excel 文档解析器"""

    def __init__(self, cleaner: DataCleaner = None, table_cleaner: TableCleaner = None):
        self.cleaner = cleaner or DataCleaner()
        self.table_cleaner = table_cleaner or TableCleaner()

    def parse(self, file_path: str, enable_cleaning: bool = True) -> ParsedDocument:
        """解析 Excel 文件"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        wb = load_workbook(file_path, data_only=True)
        pages = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = list(ws.rows)
            if not rows:
                continue

            # 转换为文本
            text_lines = []
            for row in rows:
                row_text = []
                for cell in row:
                    if cell.value is not None:
                        val = str(cell.value).strip()
                        if val:
                            row_text.append(val)
                if row_text:
                    text_lines.append(" | ".join(row_text))

            text = "\n".join(text_lines)
            if enable_cleaning:
                text = self.cleaner.clean_text(text)

            pages.append(DocumentPage(
                page_num=len(pages) + 1,
                text=text,
                tables=[],  # Excel 本身是表格，不单独提取
                metadata={'sheet_name': sheet_name}
            ))

        return ParsedDocument(
            file_path=file_path,
            file_name=os.path.basename(file_path),
            file_type='excel',
            pages=pages,
            total_pages=len(pages)
        )

    def parse_simple(self, file_path: str) -> str:
        """简单解析，只返回纯文本"""
        doc = self.parse(file_path, enable_cleaning=False)
        return "\n\n".join([page.text for page in doc.pages if page.text])


class DocumentParser:
    """统一文档解析器 - 根据文件类型自动选择解析器"""

    def __init__(self, cleaner: DataCleaner = None):
        self.cleaner = cleaner or DataCleaner()
        self.table_cleaner = TableCleaner()
        self._parsers = {
            'pdf': PDFParser(self.cleaner, self.table_cleaner),
            'docx': DOCXParser(self.cleaner),
            'txt': TXTParser(self.cleaner),
            'xlsx': ExcelParser(self.cleaner, self.table_cleaner),
            'xls': ExcelParser(self.cleaner, self.table_cleaner),
        }

    def get_file_type(self, file_path: str) -> str:
        """根据扩展名获取文件类型"""
        ext = os.path.splitext(file_path)[1].lower().lstrip('.')
        if ext in self._parsers:
            return ext
        return 'txt'  # 默认当作文本处理

    def parse(
            self,
            file_path: str,
            enable_cleaning: bool = True,
            **kwargs
    ) -> ParsedDocument:
        """
        解析文档

        Args:
            file_path: 文件路径
            enable_cleaning: 是否启用清洗
            **kwargs: 传递给具体解析器的参数

        Returns:
            ParsedDocument 对象
        """
        file_type = self.get_file_type(file_path)
        parser = self._parsers.get(file_type, self._parsers['txt'])

        return parser.parse(file_path, enable_cleaning=enable_cleaning, **kwargs)

    def parse_to_text(self, file_path: str, **kwargs) -> str:
        """解析文档并返回纯文本"""
        doc = self.parse(file_path, **kwargs)
        return "\n\n".join([page.text for page in doc.pages if page.text])

    def parse_simple(self, file_path: str) -> str:
        """简单解析（不进行清洗）"""
        file_type = self.get_file_type(file_path)
        parser = self._parsers.get(file_type, self._parsers['txt'])

        if hasattr(parser, 'parse_simple'):
            return parser.parse_simple(file_path)
        else:
            return self.parse_to_text(file_path, enable_cleaning=False)


# ==================== 便捷函数 ====================

def parse_document(file_path: str, enable_cleaning: bool = True) -> ParsedDocument:
    """快速解析文档"""
    parser = DocumentParser()
    return parser.parse(file_path, enable_cleaning=enable_cleaning)


def parse_document_to_text(file_path: str, enable_cleaning: bool = False) -> str:
    """快速解析文档为纯文本"""
    parser = DocumentParser()
    return parser.parse_to_text(file_path, enable_cleaning=enable_cleaning)


# ==================== 演示 ====================

if __name__ == "__main__":
    print("=" * 60)
    print("RAG1 - 文档读取与数据清洗模块")
    print("=" * 60)

    # 测试清洗器
    cleaner = DataCleaner()
    test_text = "  这是一个  测试文本  \n\n第二行内容\n\n   "
    cleaned = cleaner.clean_text(test_text, verbose=True)
    print(f"\n清洗结果: '{cleaned}'")

    print("\n" + "=" * 60)
    print("模块加载完成")
    print("=" * 60)