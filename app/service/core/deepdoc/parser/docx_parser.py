# app/service/core/deepdoc/parser/docx_parser.py
from io import BytesIO
import pandas as pd


class RAGFlowDocxParser:
    """DOCX 文档解析器"""

    def __call__(self, fnm, from_page=0, to_page=100000):
        """解析 DOCX 文件"""
        from docx import Document

        if isinstance(fnm, str):
            self.doc = Document(fnm)
        else:
            self.doc = Document(BytesIO(fnm))

        secs = []  # 段落内容
        for p in self.doc.paragraphs:
            if p.text.strip():
                secs.append((p.text.strip(), p.style.name))

        tbls = [self.__extract_table_content(tb) for tb in self.doc.tables]  # 提取表格
        return secs, tbls

    def __extract_table_content(self, tb):
        """提取表格内容"""
        df = []
        for row in tb.rows:
            df.append([c.text for c in row.cells])
        return self.__compose_table_content(pd.DataFrame(df))

    def __compose_table_content(self, df):
        """将表格转换为文本格式"""
        if df.empty:
            return ""

        result = []
        # 添加表头
        headers = df.columns.tolist() if isinstance(df.columns, pd.Index) else list(range(len(df.columns)))
        result.append(" | ".join(str(h) for h in headers))
        result.append(" | ".join(["---"] * len(headers)))

        # 添加数据行
        for _, row in df.iterrows():
            result.append(" | ".join(str(v) for v in row.values))

        return "\n".join(result)